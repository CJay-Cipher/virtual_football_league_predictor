from docx import Document
import openai

# Set your OpenAI API key
api_key = "YOUR API KEY"
openai.api_key = api_key

def chat_with_bot(prompt):
    # Generate a response using the OpenAI Completion API
    response = openai.Completion.create(
        engine='text-davinci-003',
        prompt=prompt,
        max_tokens=100,
        n=1,
        stop=None,
        temperature=0.7
    )

    # Extract the text of the generated response
    generated_text = response.choices[0].text.strip()

    return generated_text


def extract_professional_summary(doc):
    """
    Extracts the professional summary from the document.

    Args:
        doc (Document): The document object.

    Returns:
        str: The extracted professional summary.
    """
    professional_summary = ""
    found_summary = False

    for paragraph in doc.paragraphs:
        if found_summary:
            if paragraph.text.strip():
                professional_summary = paragraph.text
                print()
                break
        elif "professional summary" in paragraph.text.lower():
            found_summary = True

    return professional_summary.strip()


def save_modified_docx(doc, output_path):
    """
    Saves the modified document to the specified output path.

    Args:
        doc (Document): The document object.
        output_path (str): The output file path.
    """
    doc.save(output_path)
    print(f"Modified resume saved as {output_path}")


def main():
    # Input and output file paths
    input_path = "sample_resume.docx"  # input docx file path
    output_path = "sample_output.docx"  # modified docx file path

    # Load the document
    doc = Document(input_path)

    # Extract the professional summary
    professional_summary = extract_professional_summary(doc)

    # New summary to replace the existing one using api prompt
    user_input = f"give me a modified professional summary based on this; \n{professional_summary}"
    prompt = f"You: {user_input}\nBot:"
    response = chat_with_bot(prompt)  # ChatGPT response
    # print(response)
    new_summary = response

    # Modify the document with the new summary
    modify_text = False
    for paragraph in doc.paragraphs:
        if modify_text:
            if paragraph.text.strip():
                break
            else:
                modify_text = False
        elif professional_summary.startswith(paragraph.text.strip()):
            if len(paragraph.text.strip()) != 0:
                modify_text = True
                paragraph.text = new_summary

    # Save the modified document
    save_modified_docx(doc, output_path)


if __name__ == "__main__":
    main()